from flask import Flask, request, render_template
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions
from werkzeug.utils import secure_filename
from datetime import datetime, timedelta
from openpyxl import load_workbook
from sentence_transformers import SentenceTransformer
from elasticsearch import Elasticsearch, helpers
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
import os, uuid, json, docx, pdfplumber, spacy
from io import BytesIO
from dotenv import load_dotenv
from sentence_transformers import util

load_dotenv()

import unicodedata

def eliminar_acentos(texto):
    nfkd = unicodedata.normalize('NFKD', texto)
    return ''.join([c for c in nfkd if not unicodedata.combining(c)]).lower()

import olefile

def metadades_xls(path):
    try:
        ole = olefile.OleFileIO(path)
        if ole.exists('SummaryInformation'):
            meta = ole.getproperties('SummaryInformation')
            autor = meta.get(4, 'desconegut')  # ID 4 = author
            data_creacio = meta.get(12, 'desconeguda')  # ID 12 = create date
            return {"autor": autor, "data_creacio": str(data_creacio)}
    except Exception as e:
        print("Error llegint .xls:", e)
    return {"autor": "invàlid", "data_creacio": "invàlida"}

# Configuració
AZURE_CONNECTION_STRING = os.getenv("AZURE_CONNECTION_STRING")
ELASTIC_CLOUD_ID = os.getenv("ELASTIC_CLOUD_ID")
ELASTIC_API_KEY = os.getenv("ELASTIC_API_KEY")
CONTAINER_NAME = "originals"
INDEX_NAME = "documentos_a_elegir"

app = Flask(__name__)
model = SentenceTransformer('all-MiniLM-L6-v2')
nlp = spacy.load("es_core_news_sm")

# Connexió
blob_service_client = BlobServiceClient.from_connection_string(AZURE_CONNECTION_STRING)
container_client = blob_service_client.get_container_client(CONTAINER_NAME)
es = Elasticsearch(cloud_id=ELASTIC_CLOUD_ID, api_key=ELASTIC_API_KEY)

print("🔗 Connexió amb Elasticsearch:")
print("Cloud ID:", ELASTIC_CLOUD_ID)
print("API Key (parcial):", ELASTIC_API_KEY[:5] + "..." + ELASTIC_API_KEY[-5:])
def guardar_a_blob_storage(blob_data, blob_name, metadata=None):
    try:
        blob_client = container_client.get_blob_client(blob_name)
        blob_client.upload_blob(blob_data, overwrite=True, metadata=metadata)
        print(f"✅ Fitxer {blob_name} pujat correctament.")
    except Exception as e:
        print(f"❌ Error pujant fitxer a Azure: {e}")

def generar_url_sas(nom_fitxer):
    sas_token = generate_blob_sas(
        account_name="stbmartinez2etl",
        container_name=CONTAINER_NAME,
        blob_name=nom_fitxer,
        permission=BlobSasPermissions(read=True),
        expiry=datetime.utcnow() + timedelta(hours=1)
    )
    return f"https://stbmartinez2etl.blob.core.windows.net/originals/{nom_fitxer}?{sas_token}"

def metadades_docx(path):
    props = docx.Document(path).core_properties
    return {
        "creador": props.author or "Desconegut",
        "data_creacio": props.created.isoformat() if props.created else "Desconeguda",
        "tema": props.subject or "Desconegut",
    }

# def metadades_excel(path):
#     import os
#     if not os.path.exists(path):
#         print("El archivo no existe en la ruta proporcionada.")

#     props = load_workbook(path).properties
#     return {
#         "creador": props.creator or "Desconegut",
#         "data_creacio": props.created.isoformat() if props.created else "Desconeguda",
#         "tema": props.subject or "Desconegut",
#     }

def filtrar_paraules_clau(text):
    doc = nlp(text)
    return " ".join([
    token.lemma_ for token in doc
    if token.pos_ in ("NOUN", "PROPN", "VERB", "ADJ")
    ])

def crear_indice():
    if not es.indices.exists(index=INDEX_NAME):
        print(f"🧱 Creant índex nou: {INDEX_NAME}")
        es.indices.create(index=INDEX_NAME, body={
            "settings": {
                "analysis": {
                    "tokenizer": {
                        "standard": {
                            "type": "standard"
                        }
                    },
                    "filter": {
                        "asciifolding": {
                            "type": "asciifolding",
                            "preserve_original": "false"
                        }
                    },
                    "analyzer": {
                        "custom_analyzer": {
                            "type": "custom",
                            "tokenizer": "standard",
                            "filter": ["lowercase", "asciifolding"]
                        }
                    }
                }
            },
            "mappings": {
                "properties": {
                    "nombre_archivo": {"type": "keyword"},
                    "contenido": {
                        "type": "text",
                        "analyzer": "custom_analyzer"  # Usar el analizador personalizado
                    },
                    "embedding": {"type": "dense_vector", "dims": 384, "index": True, "similarity": "cosine"},
                    "autor": {"type": "text"},
                    "data_creacio": {"type": "date"},
                    "tema": {"type": "keyword"}
                }
            }
        })
    else:
        print(f"ℹ️ L'índex {INDEX_NAME} ja existeix.")

def extraer_texto_desde_blob(blob_data):
    doc = Document(BytesIO(blob_data))
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def indexar_documents():
    docs = []
    for blob in container_client.list_blobs():
        if blob.name.endswith((".docx", ".pdf", ".xlsx")):
            print(f"📄 Indexant document: {blob.name}")
            blob_client = container_client.get_blob_client(blob)
            content = blob_client.download_blob().readall()
            if blob.name.endswith(".docx"):
                metadata = metadades_docx(BytesIO(content))
            else:
                metadata = blob_client.get_blob_properties().metadata


            if blob.name.endswith(".docx"):
                text = extraer_texto_desde_blob(content)
            elif blob.name.endswith(".pdf"):
                text = extraer_texto_pdf(content)

            embedding = model.encode(text).tolist()
            docs.append({
                "_index": INDEX_NAME,
                "_id": blob.name,
                "_source": {
                    "nombre_archivo": blob.name,
                    "contenido": text,
                    "embedding": embedding,
                    "autor": metadata.get("creador"),
                    "data_creacio": metadata.get("data_creacio", None),
                    "tema": metadata.get("tema")
                }
            })

    if docs:
        print(f"✅ {len(docs)} documents indexats.")
        helpers.bulk(es, docs)
def trobar_fragment_mes_semblant(text_llarg, consulta, mida_fragment=500):
    trossos = [text_llarg[i:i + mida_fragment] for i in range(0, len(text_llarg), mida_fragment)]
    emb_consulta = model.encode(consulta, convert_to_tensor=True)
    emb_trossos = model.encode(trossos, convert_to_tensor=True)

    similituds = util.cos_sim(emb_consulta, emb_trossos)[0]
    idx_max = similituds.argmax().item()
    fragment = trossos[idx_max]
    return fragment

def executar_consulta(embedding, text, filtres, paraula=None):
    print("🔍 Executant consulta amb filtres:", json.dumps(filtres, indent=2, ensure_ascii=False))
    resultats = es.search(index=INDEX_NAME, body={
        "query": {
            "bool": {
                "must": filtres,
                "should": [
                    {
                        "knn": {
                            "field": "embedding",
                            "query_vector": embedding,
                            "k": 10,
                            "num_candidates": 100
                        }
                    },
                    {
                        "match": {
                            "contenido": text
                        }
                    }
                ],
                "minimum_should_match": 1
            }
        },
        "highlight": {
            "fields": {
                "contenido": {
                    "fragment_size": 150,
                    "number_of_fragments": 1
                }
            }
        },
        "size": 5  # Limitar los resultados a un máximo de 5 documentos
    })

    resultats_unics = {}
    for hit in resultats["hits"]["hits"]:
        doc_id = hit["_id"]
        if doc_id not in resultats_unics:
            resultats_unics[doc_id] = hit  # només el primer per document

    # Mostrar solo un resultado por documento
    resultats_filtrats = list(resultats_unics.values())

    # Mostrar solo el primer fragmento destacado si existe
    for idx, hit in enumerate(resultats_filtrats):
        print(f"\n📁 Resultat #{idx + 1}")
        print("📄 Arxiu:", hit['_source'].get("nombre_archivo", "Desconegut"))
        print("👤 Autor:", hit['_source'].get("autor", "Desconegut"))
        print("📅 Data de creació:", hit['_source'].get("data_creacio", "Desconeguda"))
        print("📝 Tema:", hit['_source'].get("tema", "Desconegut"))
        print("⭐ Score:", round(hit['_score'], 2))

        # Ahora dentro del ciclo
        highlight = hit.get("highlight", {}).get("contenido", [])
        
        if highlight:
            fragment = max(highlight, key=len)
        else:
            contingut = hit["_source"].get("contenido", "")
            if paraula:
                fragment = trobar_fragment_mes_semblant(contingut, paraula)
            else:
                paraules = contingut.split()
                fragment = " ".join(paraules[:100]) + "..." if paraules else "---"
        print("🧠 Fragment més rellevant:", fragment)

        # Guardar el fragment en el documento
        hit["_source"]["fragment"] = fragment

    # Devolver los resultados con los fragmentos destacados
    resultats_final = []
    for hit in resultats_filtrats:
        highlight = hit.get("highlight", {}).get("contenido", [])
        if highlight:
            fragment = max(highlight, key=len)
        else:
            contingut = hit["_source"].get("contenido", "")
            paraules = contingut.split()
            fragment = " ".join(paraules[:100]) + "..." if paraules else "---"

        hit["_source"]["fragment"] = fragment
        resultats_final.append(hit)

    return resultats_final

def extraer_texto_docx(path, paraules_per_pagina=300):
    text = " ".join([p.text for p in docx.Document(path).paragraphs if p.text.strip()])
    paraules = text.split()
    return [" ".join(paraules[i:i+paraules_per_pagina]) for i in range(0, len(paraules), paraules_per_pagina)]

def extraer_texto_pdf(blob_data):
    with pdfplumber.open(BytesIO(blob_data)) as pdf:
        return "\n".join([page.extract_text() or "" for page in pdf.pages])

# def extraer_texto_excel(blob_data):
#     wb = load_workbook(filename=BytesIO(blob_data), data_only=True)
#     texto = []
#     for hoja in wb.worksheets:
#         for fila in hoja.iter_rows(values_only=True):
#             fila_text = " ".join([str(celda) for celda in fila if celda is not None])
#             if fila_text.strip():
#                 texto.append(fila_text)
#     return "\n".join(texto)

def dividir_text_per_pagina(text, paraules_per_pagina=300):
    paraules = text.split()
    return [" ".join(paraules[i:i+paraules_per_pagina]) for i in range(0, len(paraules), paraules_per_pagina)]
from pypdf import PdfReader

@app.route("/")
def index():
    return render_template("index.html")
import pandas as pd
import json
import pyarrow.parquet as pq
import pyarrow as pa

import pandas as pd
import pyarrow.parquet as pq
import pyarrow as pa
from azure.storage.blob import BlobServiceClient, ContentSettings
from datetime import datetime
import os
import io

def subir_parquet_documento(documentos, conexion_str, contenedor="informacion"):
    datos = []
    nombres_orden = ["primer resultado", "segundo resultado"]

    fecha_ejecucion = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    for i, doc in enumerate(documentos):
        nombre_archivo = doc.get("nom", "")
        extension = os.path.splitext(nombre_archivo)[1].lower()

        datos.append({
            "nombre_archivo": nombre_archivo,
            "extension": extension,
            "autor": doc.get("autor"),
            "data_creacio": doc.get("data"),
            "tema": doc.get("tema"),
            "frase_filtrada": doc.get("frase_filtrada"),
            "score": doc.get("score"),
            "fragment": doc.get("fragment"),
            "numero_paginas": len(doc.get("pagines", [])),
            "orden": nombres_orden[i] if i < len(nombres_orden) else f"{i+1}º resultado",
            "fecha_ejecucion": fecha_ejecucion  # 🆕 Se añade aquí la fecha y hora
        })

    df = pd.DataFrame(datos)

    # Crear archivo Parquet en memoria
    table = pa.Table.from_pandas(df)
    buffer = io.BytesIO()
    pq.write_table(table, buffer)
    buffer.seek(0)

    # Formatem la data d'avui per crear una carpeta
    data_actual = datetime.now().strftime('%Y-%m-%d')  # ex: "2025-05-11"
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')  # ex: "20250511_150010"

    # Construïm el nom del blob amb la "carpeta" de la data
    nombre_blob = f"{data_actual}/resultats_{timestamp}.parquet"

    # Upload a Azure Blob Storage
    blob_service_client = BlobServiceClient.from_connection_string(conexion_str)
    container_client = blob_service_client.get_container_client(contenedor)
    blob_client = container_client.get_blob_client(nombre_blob)

    blob_client.upload_blob(
        buffer,
        overwrite=True,
        content_settings=ContentSettings(content_type='application/octet-stream')
    )

    print(f"✅ Parquet '{nombre_blob}' pujat correctament a {contenedor}")


from azure.storage.blob import BlobServiceClient, ContentSettings
import os

def subir_parquet_a_blob(ruta_parquet, nombre_blob, conexion_str, contenedor="informacion"):
    """
    Sube un archivo Parquet a Azure Blob Storage.
    """
    blob_service_client = BlobServiceClient.from_connection_string(conexion_str)
    container_client = blob_service_client.get_container_client(contenedor)

    # Crear contenedor si no existe
    try:
        container_client.create_container()
    except Exception:
        pass  # Ya existe

    blob_client = container_client.get_blob_client(nombre_blob)

    with open(ruta_parquet, "rb") as data:
        blob_client.upload_blob(
            data,
            overwrite=True,
            content_settings=ContentSettings(content_type="application/octet-stream")
        )

    print(f"✅ Archivo '{nombre_blob}' subido a Azure Blob Storage en contenedor '{contenedor}'")

@app.route("/carregar", methods=["POST"])
def carregar():
    paraula = request.form["paraula"]
    fitxers = request.files.getlist("fitxers")

    # 🔍 Recollir filtres del formulari per fer la cerca
    filtre_creador = request.form.get("creador")
    filtre_data = request.form.get("data")
    filtre_tema = request.form.get("tema")

    for fitxer in fitxers:
        if fitxer.filename == "":
            continue

        nom_fitxer = secure_filename(fitxer.filename)
        blob_data = fitxer.read()

        # 📦 Metadades del fitxer
        if nom_fitxer.endswith(".docx"):
            with BytesIO(blob_data) as f:
                props = docx.Document(f).core_properties
                metadata_creador = props.author or "Desconegut"
                metadata_data = props.created.isoformat() if props.created else datetime.utcnow().isoformat()
                metadata_tema = props.subject or "Desconegut"
        elif nom_fitxer.endswith(".pdf"):
            with BytesIO(blob_data) as f:
                pdf = PdfReader(f)
                props = pdf.metadata
                metadata_creador = props.get("/Author", "Desconegut")
                metadata_data = props.get("/CreationDate", datetime.utcnow().isoformat())
                if metadata_data.startswith("D:"):
                    metadata_data = metadata_data[2:10]  # Format YYYYMMDD
                    try:
                        metadata_data = datetime.strptime(metadata_data, "%Y%m%d").isoformat()
                    except:
                        metadata_data = datetime.utcnow().isoformat()
                metadata_tema = props.get("/Subject", "Desconegut")

        else:
            metadata_creador = "Desconegut"
            metadata_data = datetime.utcnow().isoformat()
            metadata_tema = "Desconegut"

        metadata = {
            "creador": metadata_creador,
            "data_creacio": metadata_data,
            "tema": metadata_tema
        }

        print(f"\n📦 Metadades del fitxer '{nom_fitxer}':")
        for k, v in metadata.items():
            print(f"  - {k}: {v}")

        blob_client = container_client.get_blob_client(nom_fitxer)
        blob_client.upload_blob(blob_data, overwrite=True, metadata=metadata)
        print(f"✅ Fitxer '{nom_fitxer}' pujat amb metadades.")

    # 🔎 Construcció de filtres per la cerca (basats en el formulari)
    filtres = []
    if filtre_creador:
        filtres.append({"match": {"autor": filtre_creador}})
    if filtre_data:
        filtres.append({"match": {"data_creacio": filtre_data}})
    if filtre_tema:
        filtres.append({"match": {"tema": filtre_tema}})

    crear_indice()
    indexar_documents()

    top_docs = []
    if paraula or filtres:
        paraula_filtrada = filtrar_paraules_clau(paraula) if paraula else ""
        paraula_filtrada = eliminar_acentos(paraula_filtrada) if paraula_filtrada else ""

        # Solo crear el embedding si hay una palabra clave válida
        if paraula_filtrada:
            embedding = model.encode(paraula_filtrada).tolist()
        else:
            embedding = None  # No generamos un embedding vacío
        
        if embedding:
            resultats_unics = executar_consulta(embedding, paraula_filtrada, filtres, paraula=paraula or "")
        else:
            # Si no hay embedding, solo aplicar los filtros sin consulta por embeddings
            resultats_unics = executar_consulta([], paraula_filtrada, filtres, paraula=paraula or "")

        for hit in resultats_unics[:2]:
            contingut = hit["_source"].get("contenido", "")
            pagines = dividir_text_per_pagina(contingut, paraules_per_pagina=300)
            top_docs.append({
                "nom": hit["_source"].get("nombre_archivo", "Desconegut"),
                "autor": hit["_source"].get("autor", "Desconegut"),
                "data": hit["_source"].get("data_creacio", "Desconegut"),
                "tema": hit["_source"].get("tema", "Desconegut"),
                "frase_filtrada": paraula_filtrada,
                "score": round(hit["_score"], 2),
                "fragment": hit["_source"].get("fragment", "---"),
                "pagines": pagines
            })
        # Guardar resultados como JSON y Parquet
            docs_to_upload = top_docs[:2]  # Selecciona hasta 2 documentos si hay disponibles
    if docs_to_upload:
        subir_parquet_documento(docs_to_upload, AZURE_CONNECTION_STRING)
    else:
        print("⚠️ No se encontró ningún documento para subir.")







    return render_template("resultat.html", documents=top_docs, paraula=paraula)

if __name__ == "__main__":
    app.run(debug=True)