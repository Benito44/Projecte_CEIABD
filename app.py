from flask import Flask, request, render_template
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions, ContentSettings
from werkzeug.utils import secure_filename
from datetime import datetime, timedelta
from sentence_transformers import SentenceTransformer, util
from elasticsearch import Elasticsearch, helpers
from docx import Document
from pypdf import PdfReader
import pdfplumber
import os, io, json, unicodedata, spacy, pandas as pd, pyarrow as pa, pyarrow.parquet as pq
from io import BytesIO
from dotenv import load_dotenv

load_dotenv()

# Config
AZURE_CONNECTION_STRING = os.getenv("AZURE_CONNECTION_STRING")
ELASTIC_CLOUD_ID = os.getenv("ELASTIC_CLOUD_ID")
ELASTIC_API_KEY = os.getenv("ELASTIC_API_KEY")
CONTAINER_NAME = "originals"
INDEX_NAME = "azure_doc"

app = Flask(__name__)
model = SentenceTransformer('all-MiniLM-L6-v2')
nlp = spacy.load("es_core_news_sm")

# Connexions
blob_service_client = BlobServiceClient.from_connection_string(AZURE_CONNECTION_STRING)
container_client = blob_service_client.get_container_client(CONTAINER_NAME)
es = Elasticsearch(cloud_id=ELASTIC_CLOUD_ID, api_key=ELASTIC_API_KEY)
print("🔗 Connexió amb Elasticsearch:")
print("Cloud ID:", ELASTIC_CLOUD_ID)
print("API Key (parcial):", ELASTIC_API_KEY[:5] + "..." + ELASTIC_API_KEY[-5:])

# Utils

def eliminar_acentos(texto):
    nfkd = unicodedata.normalize('NFKD', texto)
    return ''.join([c for c in nfkd if not unicodedata.combining(c)]).lower()

def filtrar_paraules_clau(text):
    doc = nlp(text)
    return " ".join([t.lemma_ for t in doc if t.pos_ in ("NOUN", "PROPN", "VERB", "ADJ")])

def dividir_text(text, mida=300):
    paraules = text.split()
    return [" ".join(paraules[i:i + mida]) for i in range(0, len(paraules), mida)]

def extraer_texto_docx(blob):
    return " ".join([p.text for p in Document(BytesIO(blob)).paragraphs if p.text.strip()])

def extraer_texto_pdf(blob):
    with pdfplumber.open(BytesIO(blob)) as pdf:
        return "\n".join([p.extract_text() or "" for p in pdf.pages])

def trobar_fragment(text, consulta):
    trossos = dividir_text(text, 500)
    emb_consulta = model.encode(consulta, convert_to_tensor=True)
    emb_trossos = model.encode(trossos, convert_to_tensor=True)
    similituds = util.cos_sim(emb_consulta, emb_trossos)[0]
    return trossos[similituds.argmax().item()]

def crear_indice():
    if not es.indices.exists(index=INDEX_NAME):
        print(f"🧱 Creant índex nou: {INDEX_NAME}")
        es.indices.create(index=INDEX_NAME, body={
            "settings": {"analysis": {"analyzer": {"custom_analyzer": {"type": "custom", "tokenizer": "standard", "filter": ["lowercase", "asciifolding"]}}}},
            "mappings": {"properties": {
                "nombre_archivo": {"type": "keyword"},
                "contenido": {"type": "text", "analyzer": "custom_analyzer"},
                "embedding": {"type": "dense_vector", "dims": 384, "index": True, "similarity": "cosine"},
                "autor": {"type": "text"},
                "data_creacio": {"type": "date"},
                "tema": {"type": "keyword"}
            }}
        })
    else:
        print(f"ℹ️ L'índex {INDEX_NAME} ja existeix.")

def indexar_documents():
    docs = []
    for blob in container_client.list_blobs():
        if blob.name.endswith((".docx", ".pdf")):
            print(f"📄 Indexant document: {blob.name}")
            blob_data = container_client.get_blob_client(blob).download_blob().readall()
            text = extraer_texto_docx(blob_data) if blob.name.endswith(".docx") else extraer_texto_pdf(blob_data)
            metadata = container_client.get_blob_client(blob).get_blob_properties().metadata
            embedding = model.encode(text).tolist()
            docs.append({"_index": INDEX_NAME, "_id": blob.name, "_source": {
                "nombre_archivo": blob.name,
                "contenido": text,
                "embedding": embedding,
                "autor": metadata.get("creador", "Desconegut"),
                "data_creacio": metadata.get("data_creacio"),
                "tema": metadata.get("tema", "Desconegut")
            }})
    if docs:
        print(f"✅ {len(docs)} documents indexats.")
        helpers.bulk(es, docs)

def executar_consulta(embedding, text, filtres, paraula):
    print("🔍 Executant consulta amb filtres:", json.dumps(filtres, indent=2, ensure_ascii=False))
    query = {
        "query": {
            "bool": {
                "must": filtres,
                "should": [
                    {"knn": {"field": "embedding", "query_vector": embedding, "k": 10, "num_candidates": 100}},
                    {"match": {"contenido": text}}
                ],
                "minimum_should_match": 1
            }
        },
        "highlight": {"fields": {"contenido": {"fragment_size": 150, "number_of_fragments": 1}}},
        "size": 5
    }
    resultats = es.search(index=INDEX_NAME, body=query)
    unics = {}
    for idx, hit in enumerate(resultats["hits"]["hits"]):
        doc_id = hit["_id"]
        if doc_id not in unics:
            fragment = hit.get("highlight", {}).get("contenido", [])
            if fragment:
                fragment = max(fragment, key=len)
            else:
                contingut = hit["_source"].get("contenido", "")
                fragment = trobar_fragment(contingut, paraula)
            print(f"\n📁 Resultat #{idx + 1}")
            print("📄 Arxiu:", hit['_source'].get("nombre_archivo", "Desconegut"))
            print("👤 Autor:", hit['_source'].get("autor", "Desconegut"))
            print("📅 Data de creació:", hit['_source'].get("data_creacio", "Desconeguda"))
            print("📝 Tema:", hit['_source'].get("tema", "Desconegut"))
            print("⭐ Score:", round(hit['_score'], 2))
            print("🧠 Fragment més rellevant:", fragment)
            hit["_source"]["fragment"] = fragment
            unics[doc_id] = hit
    return list(unics.values())

def subir_parquet_documento(documentos):
    dades = []
    for i, doc in enumerate(documentos):
        dades.append({
            "nombre_archivo": doc.get("nom"),
            "extension": os.path.splitext(doc.get("nom"))[1].lower(),
            "autor": doc.get("autor"),
            "data_creacio": doc.get("data"),
            "tema": doc.get("tema"),
            "frase_filtrada": doc.get("frase_filtrada"),
            "score": doc.get("score"),
            "fragment": doc.get("fragment"),
            "numero_paginas": len(doc.get("pagines", [])),
            "orden": f"{i+1}º resultat",
            "fecha_ejecucion": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        })
    table = pa.Table.from_pandas(pd.DataFrame(dades))
    buffer = io.BytesIO()
    pq.write_table(table, buffer)
    buffer.seek(0)
    nom_blob = f"{datetime.now():%Y-%m-%d}/resultats_{datetime.now():%Y%m%d_%H%M%S}.parquet"
    print(f"✅ Parquet '{nom_blob}' pujat correctament a Azure Blob Storage")
    blob_client = blob_service_client.get_blob_client(container="informacion", blob=nom_blob)
    blob_client.upload_blob(buffer, overwrite=True, content_settings=ContentSettings(content_type='application/octet-stream'))

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/carregar", methods=["POST"])
def carregar():
    paraula = request.form.get("paraula", "")
    fitxers = request.files.getlist("fitxers")
    filtres = []
    for clau in ["creador", "data", "tema"]:
        val = request.form.get(clau)
        if val:
            filtres.append({"match": {"autor" if clau == "creador" else clau: val}})
    for fitxer in fitxers:
        if not fitxer.filename:
            continue
        nom_fitxer = secure_filename(fitxer.filename)
        blob_data = fitxer.read()
        print(f"\n📦 Fitxer carregat: {nom_fitxer}")
        if nom_fitxer.endswith(".docx"):
            props = Document(BytesIO(blob_data)).core_properties
            creador = props.author or "Desconegut"
            data = props.created.isoformat() if props.created else datetime.utcnow().isoformat()
            tema = props.subject or "Desconegut"
        elif nom_fitxer.endswith(".pdf"):
            props = PdfReader(BytesIO(blob_data)).metadata
            creador = props.get("/Author", "Desconegut")
            data = props.get("/CreationDate", datetime.utcnow().isoformat())[2:10]
            try:
                data = datetime.strptime(data, "%Y%m%d").isoformat()
            except:
                data = datetime.utcnow().isoformat()
            tema = props.get("/Subject", "Desconegut")
        else:
            creador, data, tema = "Desconegut", datetime.utcnow().isoformat(), "Desconegut"
        print(f"  - Creador: {creador}\n  - Data: {data}\n  - Tema: {tema}")
        container_client.upload_blob(name=nom_fitxer, data=blob_data, overwrite=True, metadata={"creador": creador, "data_creacio": data, "tema": tema})
        print(f"✅ Fitxer '{nom_fitxer}' pujat amb metadades.")
    crear_indice()
    indexar_documents()
    top_docs = []
    if paraula or filtres:
        paraula_filtrada = eliminar_acentos(filtrar_paraules_clau(paraula))
        embedding = model.encode(paraula_filtrada).tolist() if paraula_filtrada else []
        resultats = executar_consulta(embedding, paraula_filtrada, filtres, paraula)
        for hit in resultats[:2]:
            contingut = hit["_source"].get("contenido", "")
            top_docs.append({
                "nom": hit["_source"].get("nombre_archivo", "Desconegut"),
                "autor": hit["_source"].get("autor", "Desconegut"),
                "data": hit["_source"].get("data_creacio", "Desconeguda"),
                "tema": hit["_source"].get("tema", "Desconegut"),
                "frase_filtrada": paraula_filtrada,
                "score": round(hit["_score"], 2),
                "fragment": hit["_source"].get("fragment", "---"),
                "pagines": dividir_text(contingut)
            })
    if top_docs:
        subir_parquet_documento(top_docs)
        return render_template("resultat.html", documents=top_docs, paraula=paraula)
    else:
        print("⚠️ No s'ha trobat cap document per pujar.")
        return render_template("index.html", missatge="No s'ha trobat cap resultat.")

if __name__ == "__main__":
    app.run(debug=True)

